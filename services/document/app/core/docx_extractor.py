"""
ABOUTME: DOCX and ODT document extraction with structure preservation
ABOUTME: Handles tables, styles, and formatting from word processor files
"""

from dataclasses import dataclass
from io import BytesIO
from typing import Dict, List

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.utils.logging import logger


@dataclass
class DocxSection:
    """Extracted section from DOCX"""

    text: str
    style: str
    is_heading: bool
    level: int  # Heading level (1-6) or 0 for body
    tables: List[Dict] = None


@dataclass
class DocxTable:
    """Extracted table structure"""

    rows: int
    cols: int
    data: List[List[str]]
    has_header: bool = False


class DocxExtractor:
    """
    Extract text and structure from DOCX/ODT files
    Preserves headings, paragraphs, and tables
    """

    def extract_text(self, file_content: bytes) -> str:
        """
        Extract plain text from DOCX

        Args:
            file_content: DOCX file bytes

        Returns:
            Extracted text as string
        """
        try:
            doc = Document(BytesIO(file_content))

            # Extract all paragraph text
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Join with double newlines to preserve paragraph structure
            full_text = "\n\n".join(paragraphs)

            logger.info(f"Extracted {len(paragraphs)} paragraphs from DOCX")
            return full_text

        except Exception as e:
            logger.error(f"DOCX text extraction failed: {str(e)}")
            return ""

    def extract_structured(self, file_content: bytes) -> List[DocxSection]:
        """
        Extract text with structure preservation (headings, styles)

        Args:
            file_content: DOCX file bytes

        Returns:
            List of DocxSection objects
        """
        try:
            doc = Document(BytesIO(file_content))
            sections = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"

                # Detect heading level
                is_heading = style_name.startswith("Heading")
                level = 0
                if is_heading:
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1

                sections.append(
                    DocxSection(text=text, style=style_name, is_heading=is_heading, level=level)
                )

            logger.info(f"Extracted {len(sections)} structured sections")
            return sections

        except Exception as e:
            logger.error(f"Structured DOCX extraction failed: {str(e)}")
            return []

    def extract_tables(self, file_content: bytes) -> List[DocxTable]:
        """
        Extract tables from DOCX

        Args:
            file_content: DOCX file bytes

        Returns:
            List of DocxTable objects
        """
        try:
            doc = Document(BytesIO(file_content))
            tables = []

            for table in doc.tables:
                rows = len(table.rows)
                cols = len(table.columns) if table.rows else 0

                # Extract cell data
                data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    data.append(row_data)

                # Heuristic: first row is header if it's styled differently
                # or contains fewer numbers than subsequent rows
                has_header = False
                if len(data) > 1:
                    first_row_nums = sum(
                        1 for cell in data[0] if cell.replace(".", "").replace(",", "").isdigit()
                    )
                    second_row_nums = sum(
                        1 for cell in data[1] if cell.replace(".", "").replace(",", "").isdigit()
                    )
                    has_header = first_row_nums < second_row_nums

                tables.append(DocxTable(rows=rows, cols=cols, data=data, has_header=has_header))

            logger.info(f"Extracted {len(tables)} tables")
            return tables

        except Exception as e:
            logger.error(f"Table extraction failed: {str(e)}")
            return []

    def extract_complete(self, file_content: bytes) -> Dict:
        """
        Complete extraction with text, structure, and tables

        Args:
            file_content: DOCX file bytes

        Returns:
            Dict with all extracted data
        """
        try:
            doc = Document(BytesIO(file_content))

            sections = []

            # Iterate through document body elements
            for element in doc.element.body:
                # Handle paragraphs
                if element.tag.endswith("p"):
                    para = Paragraph(element, doc)
                    text = para.text.strip()
                    if text:
                        style_name = para.style.name if para.style else "Normal"
                        is_heading = style_name.startswith("Heading")
                        level = 0
                        if is_heading:
                            try:
                                level = int(style_name.split()[-1])
                            except (ValueError, IndexError):
                                level = 1

                        sections.append(
                            {
                                "type": "paragraph",
                                "text": text,
                                "style": style_name,
                                "is_heading": is_heading,
                                "level": level,
                            }
                        )

                # Handle tables
                elif element.tag.endswith("tbl"):
                    table = Table(element, doc)
                    rows = len(table.rows)
                    cols = len(table.columns) if table.rows else 0

                    data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        data.append(row_data)

                    sections.append({"type": "table", "rows": rows, "cols": cols, "data": data})

            # Extract metadata
            core_properties = doc.core_properties
            metadata = {
                "author": core_properties.author,
                "title": core_properties.title,
                "subject": core_properties.subject,
                "created": core_properties.created.isoformat() if core_properties.created else None,
                "modified": (
                    core_properties.modified.isoformat() if core_properties.modified else None
                ),
            }

            # Clean None values
            metadata = {k: v for k, v in metadata.items() if v is not None}

            logger.info(f"Complete extraction: {len(sections)} elements")

            return {
                "sections": sections,
                "metadata": metadata,
                "total_sections": len(sections),
                "paragraph_count": sum(1 for s in sections if s["type"] == "paragraph"),
                "table_count": sum(1 for s in sections if s["type"] == "table"),
            }

        except Exception as e:
            logger.error(f"Complete DOCX extraction failed: {str(e)}")
            return {"sections": [], "metadata": {}, "error": str(e)}

    def docx_to_markdown(self, file_content: bytes) -> str:
        """
        Convert DOCX to markdown format

        Args:
            file_content: DOCX file bytes

        Returns:
            Markdown-formatted text
        """
        try:
            doc = Document(BytesIO(file_content))
            markdown_lines = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"

                # Convert headings to markdown
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.split()[-1])
                        markdown_lines.append(f"{'#' * level} {text}")
                    except (ValueError, IndexError):
                        markdown_lines.append(f"# {text}")
                else:
                    markdown_lines.append(text)

            # Handle tables
            for table in doc.tables:
                # Create markdown table
                if table.rows:
                    # Header row
                    header = [cell.text.strip() for cell in table.rows[0].cells]
                    markdown_lines.append("| " + " | ".join(header) + " |")
                    markdown_lines.append("| " + " | ".join(["---"] * len(header)) + " |")

                    # Data rows
                    for row in table.rows[1:]:
                        cells = [cell.text.strip() for cell in row.cells]
                        markdown_lines.append("| " + " | ".join(cells) + " |")

            return "\n\n".join(markdown_lines)

        except Exception as e:
            logger.error(f"DOCX to markdown conversion failed: {str(e)}")
            return ""


# Global instance
docx_extractor = DocxExtractor()
